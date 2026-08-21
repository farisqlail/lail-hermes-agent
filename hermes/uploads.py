"""Images and documents the operator hands to a chat turn.

Deliberately short-lived. The model looks once, answers, and the file is
deleted — so nothing is stored that a later turn would have to re-send, and
the prompt never carries the same attachment twice. `hermes/cleanup.py` is the
safety net for the case this path cannot cover: a process that dies between the
upload and the answer.

Image type is decided by the leading bytes, never by the filename. An `.png`
that is really an SVG would be served back into the dashboard's own origin,
where its embedded script runs with the operator's session — so the allowed
formats are the raster ones the vision models accept, and everything else is
refused.

Documents are different: they are never served back to a browser, only read
server-side and turned into plain text, so trusting the extension the client
sent is safe — worst case a mislabelled file fails to parse.
"""
from __future__ import annotations
import base64, re, uuid
from pathlib import Path

#: Per file. A phone photo is 2-5 MB; the base64 the model call carries is a
#: third larger again, and that rides in the prompt.
MAX_UPLOAD_BYTES = 10 * 1024 * 1024


class UnsupportedImage(ValueError):
    """The bytes are not one of the accepted raster formats."""


# (leading bytes, extension, mime). WebP is checked separately: its signature
# is split across the RIFF header.
_SIGNATURES = (
    (b"\x89PNG\r\n\x1a\n", "png", "image/png"),
    (b"\xff\xd8\xff", "jpg", "image/jpeg"),
    (b"GIF87a", "gif", "image/gif"),
    (b"GIF89a", "gif", "image/gif"),
)


def sniff(data: bytes) -> tuple[str, str]:
    """(extension, mime) for supported image bytes, else raise.

    Content sniffing, not extension trust: the filename comes from the browser
    and decides nothing.
    """
    for magic, ext, mime in _SIGNATURES:
        if data.startswith(magic):
            return ext, mime
    if data[:4] == b"RIFF" and data[8:12] == b"WEBP":
        return "webp", "image/webp"
    raise UnsupportedImage(
        "Format tidak didukung. Kirim PNG, JPEG, GIF, atau WebP — "
        "SVG dan berkas non-gambar ditolak.")


def _conv_dir(base: Path, conv_id: str) -> Path | None:
    """`base/conv_id`, or None when conv_id tries to leave `base`.

    Same guard, same reason as `cleanup.purge`: conv_id arrives from the
    client, so it is never joined blindly.
    """
    if not conv_id or conv_id in (".", ".."):
        return None
    try:
        resolved = (base / conv_id).resolve()
        resolved.relative_to(base.resolve())
    except (OSError, ValueError):
        return None
    return resolved


def save(base: Path, conv_id: str, data: bytes) -> tuple[str, str]:
    """Write one image, returning (stored name, mime).

    The stored name is generated here, never taken from the client: a browser
    filename can collide, can carry a path, and is the one field an attacker
    fully controls.
    """
    if len(data) > MAX_UPLOAD_BYTES:
        raise ValueError(
            f"Gambar terlalu besar (maksimal {MAX_UPLOAD_BYTES // (1024 * 1024)} MB)")
    ext, mime = sniff(data)
    d = _conv_dir(base, conv_id)
    if d is None:
        raise ValueError("id percakapan tidak sah")
    d.mkdir(parents=True, exist_ok=True)
    name = f"{uuid.uuid4().hex}.{ext}"
    (d / name).write_bytes(data)
    return name, mime


def resolve(base: Path, conv_id: str, name: str) -> Path | None:
    """The stored file for a name the client sent back, or None.

    None covers every way this can be wrong at once — traversal, a file from
    another conversation, one already discarded — because the caller treats
    them identically: the image is simply not attached.
    """
    d = _conv_dir(base, conv_id)
    if d is None or not name or "/" in name or "\\" in name:
        return None
    p = d / name
    return p if p.is_file() else None


def data_url(path: Path) -> str:
    """The inline form the chat completions API takes. Providers cannot reach
    a path on this machine, so the bytes travel in the request."""
    _, mime = sniff(path.read_bytes()[:16])
    return f"data:{mime};base64," + base64.b64encode(path.read_bytes()).decode()


def as_content_parts(text: str, images: list[Path],
                     documents: list[Path] | None = None) -> list[dict]:
    """One user message carrying text, extracted document text, and images,
    in OpenAI's parts form. Documents ride inside the text part — they're
    read once server-side, there's no vision model involved."""
    for doc in documents or []:
        try:
            extracted = extract_text(doc)
        except Exception as e:
            # Not just UnsupportedDocument: a corrupt PDF/DOCX/XLSX raises its
            # library's own error (PdfReadError, BadZipFile, ...) — any of
            # those must degrade to an inline note, never crash the turn.
            extracted = f"(gagal dibaca: {e})"
        text += f"\n\n--- Berkas: {display_name(doc.name)} ---\n{extracted}"
    parts: list[dict] = [{"type": "text", "text": text}]
    parts += [{"type": "image_url", "image_url": {"url": data_url(p)}}
              for p in images]
    return parts


class UnsupportedDocument(ValueError):
    """The extension is not one of the accepted document formats, or the
    library needed to read it isn't installed."""


#: Read as plain text — no library needed, so these work in every install.
_TEXT_EXTENSIONS = {
    "txt", "md", "markdown", "csv", "tsv", "json", "log", "yaml", "yml",
    "xml", "ini", "toml", "py", "js", "ts", "tsx", "jsx", "html", "css",
    "sh", "java", "c", "cpp", "h", "go", "rs", "rb", "php", "sql",
}
#: Need `pip install -e .[docs]` (pypdf / python-docx / openpyxl).
_PDF_EXTENSIONS = {"pdf"}
_DOCX_EXTENSIONS = {"docx"}
_XLSX_EXTENSIONS = {"xlsx"}
_DOCUMENT_EXTENSIONS = (_TEXT_EXTENSIONS | _PDF_EXTENSIONS
                        | _DOCX_EXTENSIONS | _XLSX_EXTENSIONS)

#: Kept from one extracted document, so a large attachment can't push the
#: rest of the prompt out of the context window.
MAX_DOCUMENT_CHARS = 50_000


def doc_ext(filename: str) -> str:
    """The lowercased extension, if it's one this module knows how to read."""
    ext = Path(filename or "").suffix.lstrip(".").lower()
    if ext not in _DOCUMENT_EXTENSIONS:
        raise UnsupportedDocument(
            "Format berkas tidak didukung. Didukung: teks/kode, PDF, DOCX, XLSX.")
    return ext


def _safe_stem(filename: str) -> str:
    """Original filename, made safe to embed in a generated path. Display-only
    — never trusted for path resolution, which still runs through `resolve`."""
    stem = re.sub(r"[^A-Za-z0-9_.-]+", "_", Path(filename or "").stem)[:60]
    return stem or "berkas"


def save_document(base: Path, conv_id: str, filename: str, data: bytes) -> tuple[str, str]:
    """Write one document, returning (stored name, extension)."""
    if len(data) > MAX_UPLOAD_BYTES:
        raise ValueError(
            f"Berkas terlalu besar (maksimal {MAX_UPLOAD_BYTES // (1024 * 1024)} MB)")
    ext = doc_ext(filename)
    d = _conv_dir(base, conv_id)
    if d is None:
        raise ValueError("id percakapan tidak sah")
    d.mkdir(parents=True, exist_ok=True)
    name = f"{uuid.uuid4().hex}_{_safe_stem(filename)}.{ext}"
    (d / name).write_bytes(data)
    return name, ext


def display_name(stored_name: str) -> str:
    """The original filename `save_document` folded into the stored name."""
    uuid_part, sep, rest = stored_name.partition("_")
    return rest if sep and len(uuid_part) == 32 else stored_name


def _extract_pdf(path: Path) -> str:
    try:
        from pypdf import PdfReader
    except ImportError:
        raise UnsupportedDocument(
            "Pembaca PDF belum terpasang. Jalankan: pip install -e .[docs]")
    reader = PdfReader(str(path))
    return "\n".join(page.extract_text() or "" for page in reader.pages)


def _extract_docx(path: Path) -> str:
    try:
        import docx
    except ImportError:
        raise UnsupportedDocument(
            "Pembaca DOCX belum terpasang. Jalankan: pip install -e .[docs]")
    document = docx.Document(str(path))
    return "\n".join(p.text for p in document.paragraphs)


def _extract_xlsx(path: Path) -> str:
    try:
        import openpyxl
    except ImportError:
        raise UnsupportedDocument(
            "Pembaca XLSX belum terpasang. Jalankan: pip install -e .[docs]")
    wb = openpyxl.load_workbook(str(path), data_only=True, read_only=True)
    try:
        lines = []
        for ws in wb.worksheets:
            lines.append(f"# {ws.title}")
            for row in ws.iter_rows(values_only=True):
                lines.append(", ".join("" if v is None else str(v) for v in row))
        return "\n".join(lines)
    finally:
        # read_only keeps the zip handle open until closed — on Windows that
        # blocks the delete `discard()` does right after this turn.
        wb.close()


def extract_text(path: Path) -> str:
    """Plain text pulled from one attached document, truncated to
    `MAX_DOCUMENT_CHARS`."""
    ext = path.suffix.lstrip(".").lower()
    if ext in _TEXT_EXTENSIONS:
        text = path.read_text(encoding="utf-8", errors="replace")
    elif ext in _PDF_EXTENSIONS:
        text = _extract_pdf(path)
    elif ext in _DOCX_EXTENSIONS:
        text = _extract_docx(path)
    elif ext in _XLSX_EXTENSIONS:
        text = _extract_xlsx(path)
    else:
        raise UnsupportedDocument(f"Tidak tahu cara membaca .{ext}")
    if len(text) > MAX_DOCUMENT_CHARS:
        text = text[:MAX_DOCUMENT_CHARS] + "\n...[dipotong]"
    return text


def discard(paths: list[Path]) -> None:
    """Delete images the turn has finished with. Never raises: the answer is
    already on its way to the operator, and a file that outlives its turn is
    swept at the next start."""
    for p in paths:
        try:
            p.unlink(missing_ok=True)
        except OSError as e:
            print(f"Could not discard upload {p}: {e}")
